"""
Document ingestion endpoint — upload, chunk, embed, store (protected).

Supported formats: .txt, .md, .csv, .pdf, .docx
"""

from typing import Optional

from fastapi import APIRouter, Depends, File, Form, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.dependencies import CurrentUser, require_role
from app.models.database import Document
from app.services.ingestion import ingestion_service

router = APIRouter()

# Allowed intent categories
VALID_INTENTS = {"factual", "person", "time", "location", "explanation", "other"}

# Supported file types
SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx"}


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    import fitz  # pymupdf

    text_parts = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    import io
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


@router.post("")
async def ingest_document(
    file: UploadFile = File(...),
    intent_category: str = Form(...),
    source: Optional[str] = Form(None),
    user: CurrentUser = Depends(require_role("admin", "writer")),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a document, chunk it, generate embeddings, and store in Qdrant.

    Supported formats: .txt, .md, .csv, .pdf, .docx
    - Chunks stored in the intent_category collection
    - Each chunk carries tenant_id as metadata for isolation
    - Document record saved to PostgreSQL for tracking

    Auth: admin or writer only.
    """
    try:
        # Validate intent category
        if intent_category not in VALID_INTENTS:
            return {
                "success": False,
                "data": None,
                "error": f"Invalid intent_category '{intent_category}'. Must be one of: {sorted(VALID_INTENTS)}",
            }

        # Check file extension
        filename = file.filename or "unknown.txt"
        ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in SUPPORTED_EXTENSIONS:
            return {
                "success": False,
                "data": None,
                "error": f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}",
            }

        # Read file content
        content = await file.read()

        # Extract text based on file type
        if ext == ".pdf":
            try:
                text_content = extract_text_from_pdf(content)
            except Exception as e:
                return {
                    "success": False,
                    "data": None,
                    "error": f"Failed to parse PDF: {str(e)}",
                }
        elif ext == ".docx":
            try:
                text_content = extract_text_from_docx(content)
            except ImportError:
                return {
                    "success": False,
                    "data": None,
                    "error": "python-docx package not installed. Run: pip install python-docx",
                }
            except Exception as e:
                return {
                    "success": False,
                    "data": None,
                    "error": f"Failed to parse DOCX: {str(e)}",
                }
        else:
            # Plain text files (.txt, .md, .csv)
            try:
                text_content = content.decode("utf-8")
            except UnicodeDecodeError:
                return {
                    "success": False,
                    "data": None,
                    "error": "File must be UTF-8 encoded text",
                }

        if not text_content.strip():
            return {
                "success": False,
                "data": None,
                "error": "File is empty or contains no extractable text",
            }

        # Save document record to PostgreSQL
        doc = Document(
            tenant_id=user.tenant_id,
            intent_category=intent_category,
            content=text_content[:5000],  # Store first 5k chars for reference
            source=source or filename,
        )
        db.add(doc)
        await db.flush()

        # Ingest into vector store
        result = await ingestion_service.ingest(
            content=text_content,
            intent_category=intent_category,
            tenant_id=user.tenant_id,
            source=source or filename,
            doc_id=str(doc.doc_id),
        )

        # Update document with embedding reference
        doc.embedding_id = ",".join(result["node_ids"][:5])
        await db.commit()

        return {
            "success": True,
            "data": {
                "doc_id": str(doc.doc_id),
                "tenant_id": user.tenant_id,
                "intent_category": intent_category,
                "filename": filename,
                "file_type": ext,
                "chunks_created": result["chunks_created"],
                "text_length": len(text_content),
                "source": source or filename,
            },
            "error": None,
        }

    except Exception as e:
        await db.rollback()
        return {
            "success": False,
            "data": None,
            "error": f"Ingestion failed: {str(e)}",
        }
